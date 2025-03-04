"""
Modified chatbot.py for API integration.
This module exposes a run(input_data) function that expects a dictionary with:
    - "message": the user's query
and returns the generated chat response as a string.
"""

try:
    import pysqlite3 as sqlite3
except ImportError:
    import sqlite3

import os
from pathlib import Path
import openai
from sentence_transformers import SentenceTransformer, util
from PyPDF2 import PdfReader
import docx  # pip install python-docx
from bs4 import BeautifulSoup  # pip install beautifulsoup4

# LangChain imports
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import WebBaseLoader
from langchain_openai import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain.retrievers.self_query.base import SelfQueryRetriever
from langchain.chains.query_constructor.base import AttributeInfo

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Set OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")
if openai.api_key is None:
    raise ValueError("OpenAI API Key not set. Please set OPENAI_API_KEY in your environment.")

# Disable tokenizer parallelism warning
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# --- FILE PROCESSING FUNCTIONS ---
def load_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as file:
        reader = PdfReader(file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + " "
    return text

def load_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_text_from_html(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        soup = BeautifulSoup(file, "html.parser")
        return soup.get_text()

def clean_text(text):
    return text.replace("\n", " ").replace("  ", " ").strip()

def load_course_materials(path):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=100)
    documents = []

    if os.path.isfile(path):
        if path.endswith(".pdf"):
            text = load_text_from_pdf(path)
        elif path.endswith(".docx"):
            text = load_text_from_docx(path)
        elif path.endswith(".html"):
            text = load_text_from_html(path)
        else:
            text = ""
        if text:
            chunks = text_splitter.split_text(clean_text(text))
            for chunk in chunks:
                documents.append(Document(page_content=chunk, metadata={"source": path}))
    else:
        for root, _, files in os.walk(path):
            for file in files:
                file_path = os.path.join(root, file)
                if file.endswith(".pdf"):
                    text = load_text_from_pdf(file_path)
                elif file.endswith(".docx"):
                    text = load_text_from_docx(file_path)
                elif file.endswith(".html"):
                    text = load_text_from_html(file_path)
                else:
                    continue
                chunks = text_splitter.split_text(clean_text(text))
                for chunk in chunks:
                    documents.append(Document(page_content=chunk, metadata={"source": file_path}))
    
    return documents

# --- VECTOR STORE SETUP ---
COURSE_MATERIALS_PATH = "Mark Allen Weiss - Data structures and algorithm analysis in Java-Pearson (2012).pdf"
documents = load_course_materials(COURSE_MATERIALS_PATH)

if not documents:
    raise ValueError("No documents were extracted from the provided course materials!")

# Load embeddings
embeddings = HuggingFaceEmbeddings(model_name="all-mpnet-base-v2")

# --- FIX: Use ChromaDB in-memory mode to avoid SQLite issues ---
try:
    vectorstore = Chroma.from_documents(documents, embeddings, persist_directory=None)  # In-memory mode
except Exception as e:
    print("❌ ChromaDB failed to initialize. Switching to FAISS.")
    from langchain.vectorstores import FAISS
    vectorstore = FAISS.from_documents(documents, embeddings)  # FAISS as backup

# Metadata for self-query retriever
metadata_field_info = [
    {"name": "source", "description": "The source file of the retrieved document.", "type": "string"},
]

# Self-query retriever
self_query_retriever = SelfQueryRetriever.from_llm(
    llm=ChatOpenAI(model_name="gpt-4-turbo", temperature=0.5),
    vectorstore=vectorstore,
    document_contents="A collection of course materials relevant to the user's query.",
    metadata_field_info=metadata_field_info,
)

retriever = vectorstore.as_retriever(search_type="mmr", search_kwargs={"k": 3})

retrieval_chain = RetrievalQA.from_chain_type(
    llm=ChatOpenAI(model_name="gpt-4-turbo", temperature=0.5),
    retriever=retriever,
    chain_type="stuff",
    return_source_documents=True
)

# Load pre-trained model for semantic similarity
model = SentenceTransformer('all-MiniLM-L6-v2')

def run(input_data):
    """
    Expects input_data to be a dictionary with:
      - "message": the user's chat input.
    Returns the chatbot's response as a string.
    """
    user_message = input_data.get("message")
    if not user_message:
        return "Error: No message provided."

    # Retrieve relevant documents
    try:
        self_query_results = self_query_retriever.get_relevant_documents(user_message)
        self_query_context = " ".join([doc.page_content for doc in self_query_results])
    except Exception:
        self_query_context = ""

    try:
        retrieved_docs = retriever.get_relevant_documents(user_message)
        retrieved_context = " ".join([doc.page_content for doc in retrieved_docs])
    except Exception:
        retrieved_context = ""

    combined_context = f"{retrieved_context} {self_query_context}".strip()

    # Relevance scoring using semantic similarity
    user_embedding = model.encode(user_message, convert_to_tensor=True)
    context_embedding = model.encode(combined_context, convert_to_tensor=True)
    similarity = util.pytorch_cos_sim(user_embedding, context_embedding).item()
    relevance_score = similarity

    final_context = "No additional context is needed for this response." if relevance_score < 0.5 else combined_context

    response_prompt = f"""
You are MedBot AI, a professional chatbot assistant for medical students.

User Query:
{user_message}

Relevant Context:
{final_context}

Please provide a clear and well-structured answer.
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": response_prompt}],
            temperature=0.5,
            max_tokens=500
        )
        return response["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"Error generating response: {str(e)}"

# For testing
if __name__ == '__main__':
    sample_input = {"message": "What are the symptoms of diabetes?"}
    print("Chatbot Response:\n", run(sample_input))
