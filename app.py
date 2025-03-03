from flask import Flask, request, jsonify, Response, render_template
from flask_cors import CORS
from dotenv import load_dotenv
import openai
import os
from pathlib import Path
from sentence_transformers import SentenceTransformer, util
from PyPDF2 import PdfReader
import docx  # Ensure python-docx is installed: pip install python-docx
from bs4 import BeautifulSoup  # Ensure beautifulsoup4 is installed: pip install beautifulsoup4

# ✅ Corrected LangChain Imports
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import WebBaseLoader
from langchain_openai import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain.retrievers.self_query.base import SelfQueryRetriever
from langchain.chains.query_constructor.base import AttributeInfo

# Load environment variables
load_dotenv()

# Disable tokenizer parallelism warning
os.environ["TOKENIZERS_PARALLELISM"] = "false"

app = Flask(__name__)
CORS(app)

# Load OpenAI API Key
openai.api_key = os.getenv("OPENAI_API_KEY")
if openai.api_key is None:
    raise ValueError("⚠️ ERROR: OpenAI API Key not set. Run `export OPENAI_API_KEY='your-key-here'` in your terminal.")

# Dataset Configuration
COURSE_MATERIALS_PATH = "path/to/course/materials"  # Replace with the actual path

# ---------------------- DATA COLLECTION AND PREPROCESSING ----------------------

def load_text_from_pdf(file_path):
    """Load text from a PDF file."""
    text = ""
    with open(file_path, "rb") as file:
        reader = PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def load_text_from_docx(file_path):
    """Load text from a DOCX file."""
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_text_from_html(file_path):
    """Load text from an HTML file."""
    with open(file_path, "r", encoding="utf-8") as file:
        soup = BeautifulSoup(file, "html.parser")
        return soup.get_text()

def clean_text(text):
    """Cleans the extracted text by removing extra spaces and newlines."""
    return text.replace("\n", " ").replace("  ", " ").strip()

def load_course_materials(path):
    """Load and preprocess course materials from the specified path."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=100)
    documents = []
    
    for root, _, files in os.walk(path):
        for file in files:
            file_path = os.path.join(root, file)
            if file.endswith(".pdf"):
                text = load_text_from_pdf(file_path)
            elif file.endswith(".docx"):
                text = load_text_from_docx(file_path)
            elif file.endswith(".html"):
                text = load_text_from_html(file_path)
            else:
                continue
            
            chunks = text_splitter.split_text(clean_text(text))
            for chunk in chunks:
                documents.append(Document(page_content=chunk, metadata={"source": file_path}))
    
    return documents

# Load and process course materials
documents = load_course_materials(COURSE_MATERIALS_PATH)
if not documents:
    raise ValueError("❌ No documents were extracted from the provided course materials!")

# ---------------------- RAG SETUP ----------------------

# Setup Vector Store
embeddings = HuggingFaceEmbeddings(model_name="all-mpnet-base-v2")
vectorstore = Chroma.from_documents(documents, embeddings)

# Setup Self-Query Retriever
metadata_field_info = [
    AttributeInfo(name="source", description="The source file of the retrieved document.", type="string"),
]
self_query_retriever = SelfQueryRetriever.from_llm(
    llm=ChatOpenAI(model_name="gpt-4-turbo", temperature=0.5),
    vectorstore=vectorstore,
    document_contents="A collection of course materials relevant to the user's query.",
    metadata_field_info=metadata_field_info,
)

# Setup Advanced RAG Chain
retriever = vectorstore.as_retriever(search_type="mmr", search_kwargs={"k": 3})
retrieval_chain = RetrievalQA.from_chain_type(
    llm=ChatOpenAI(model_name="gpt-4-turbo", temperature=0.5),
    retriever=retriever,
    chain_type="stuff",
    return_source_documents=True
)

# ---------------------- FLASK ROUTES ----------------------

@app.route('/')
def home():
    """Render the chatbot UI."""
    return render_template('index.html')

# Load pre-trained model for semantic similarity
model = SentenceTransformer('all-MiniLM-L6-v2')

@app.route('/chat', methods=['POST'])
def chat():
    """Handles user chat input with advanced RAG and self-query retrieval."""
    try:
        user_message = request.json.get('message')
        if not user_message:
            return jsonify({"error": "No message provided"}), 400

        def generate_response():
            """Retrieve relevant documents, apply relevance scoring, and generate AI response."""
            try:
                # ✅ **Step 1: Retrieve Documents Using Self-Query**
                self_query_results = self_query_retriever.get_relevant_documents(user_message)
                self_query_context = " ".join([doc.page_content for doc in self_query_results])

                # ✅ **Step 2: Standard Retrieval (For Comparison)**
                retrieved_docs = retriever.get_relevant_documents(user_message)
                retrieved_context = " ".join([doc.page_content for doc in retrieved_docs])

                # ✅ **Step 3: Merge Retrieved Contexts**
                combined_context = f"{retrieved_context} {self_query_context}".strip()

                # ✅ **Step 4: Relevance Scoring Using Semantic Similarity**
                user_embedding = model.encode(user_message, convert_to_tensor=True)
                context_embedding = model.encode(combined_context, convert_to_tensor=True)
                similarity = util.pytorch_cos_sim(user_embedding, context_embedding).item()

                relevance_score = similarity  # Use similarity as the relevance score

                # ✅ **Step 5: Use Context Only if Relevant**
                if relevance_score < 0.5:
                    final_context = "No additional context is needed for this response."
                else:
                    final_context = combined_context

                # ✅ **Step 6: Generate AI Response**
                response_prompt = f"""
                You are MedBot AI, a professional chatbot assistant for medical students.

                **User Query:**
                {user_message}

                **Relevant Context:**
                {final_context}

                ---
                ### **Instructions for AI:**
                - Write in a **clean, structured, and natural tone** like ChatGPT.
                - **DO NOT** add extra spaces before punctuation.
                - **DO NOT** break words unnaturally.
                - Use proper paragraph spacing and **avoid bullet points if unnecessary**.
                - If the question is casual, respond in a conversational tone.

                ### **Response:**
                """
                client = openai.OpenAI()
                response = client.chat.completions.create(
                    model="gpt-4-turbo",
                    messages=[{"role": "user", "content": response_prompt}],
                    stream=True
                )

                # ✅ **Step 7: Stream Response Token by Token**
                for chunk in response:
                    if hasattr(chunk.choices[0].delta, "content") and chunk.choices[0].delta.content is not None:
                        yield (chunk.choices[0].delta.content + "\n").encode("utf-8")

            except Exception as e:
                yield f"Error generating response: {str(e)}".encode("utf-8")

        return Response(generate_response(), content_type="text/event-stream")

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    

@app.route('/speak', methods=['POST'])
def speak():
    """Convert chatbot text response to speech with user-selected voice and speed."""
    try:
        data = request.json
        text = data.get('text')
        voice = data.get('voice', 'alloy')  # Default to 'alloy' if none selected
        speed = data.get('speed', 1.0)  # Default speed is 1.0

        if not text:
            return jsonify({"error": "No text provided"}), 400

        if voice not in ["alloy", "echo", "fable", "onyx", "nova", "shimmer"]:
            return jsonify({"error": "Invalid voice selected"}), 400

        if not (0.25 <= speed <= 4.0):
            return jsonify({"error": "Invalid speed value. Must be between 0.25 and 4.0"}), 400

        # Define the path to save the audio file
        speech_file_path = Path("static/speech.mp3")

        # Create OpenAI client
        client = openai.OpenAI()

        # Generate speech from text
        response = client.audio.speech.create(
            model="tts-1-hd",
            voice=voice,
            input=text,
            speed=speed
        )

        # Stream response to a file
        with open(speech_file_path, 'wb') as f:
            for chunk in response.iter_bytes():
                f.write(chunk)

        # Return the path to the generated audio file
        return jsonify({"audio_url": f"/{speech_file_path}"}), 200

    except Exception as e:
        print(f"❌ ERROR in /speak: {str(e)}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)